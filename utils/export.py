import io
import re
from fpdf import FPDF
from docx import Document


def _parse_lines(text: str):
    """Yield (kind, content) for each line of markdown text."""
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            yield "blank", ""
            continue
        m = re.match(r"^(#{1,6})\s+(.+)", stripped)
        if m:
            yield "heading", (len(m.group(1)), m.group(2))
            continue
        if re.match(r"^[-*+]\s+", stripped):
            content = re.sub(r"^[-*+]\s+", "", stripped)
            content = re.sub(r"\*\*(.+?)\*\*", r"\1", content)
            content = re.sub(r"\*(.+?)\*", r"\1", content)
            yield "bullet", content
            continue
        content = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
        content = re.sub(r"\*(.+?)\*", r"\1", content)
        content = re.sub(r"`(.+?)`", r"\1", content)
        yield "body", content


def _latin1(text: str) -> str:
    return text.encode("latin-1", errors="replace").decode("latin-1")


def to_pdf(text: str, doc_type: str = "") -> bytes:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    if doc_type:
        pdf.set_font("Helvetica", style="B", size=18)
        pdf.multi_cell(0, 10, _latin1(doc_type))
        pdf.ln(4)

    for kind, content in _parse_lines(text):
        if kind == "blank":
            pdf.ln(3)
        elif kind == "heading":
            level, title = content
            pdf.set_font("Helvetica", style="B", size=max(11, 16 - level * 2))
            pdf.multi_cell(0, 7, _latin1(title))
            pdf.ln(1)
        elif kind == "bullet":
            pdf.set_font("Helvetica", size=11)
            pdf.multi_cell(0, 6, _latin1(f"  •  {content}"))
        else:
            pdf.set_font("Helvetica", size=11)
            pdf.multi_cell(0, 6, _latin1(content))

    return bytes(pdf.output())


def to_docx(text: str, doc_type: str = "") -> bytes:
    doc = Document()

    if doc_type:
        doc.add_heading(doc_type, 0)

    for kind, content in _parse_lines(text):
        if kind == "blank":
            continue
        elif kind == "heading":
            level, title = content
            doc.add_heading(title, level=min(level, 3))
        elif kind == "bullet":
            doc.add_paragraph(content, style="List Bullet")
        else:
            doc.add_paragraph(content)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
